"""
DOCXResumeParser

Uses python-docx for structured extraction.

DOCX files have richer structure than PDFs:
  - Paragraph styles (Heading 1/2/3, Normal, List Bullet)
  - Tables (sometimes used for two-column layouts)
  - Inline formatting (bold often marks section headers or job titles)

We leverage this structure for higher-confidence section detection.
"""

from __future__ import annotations

from app.core.exceptions import ParseError
from app.core.logging import get_logger
from app.schemas.resume import ParsedResume
from app.parsers.base import BaseResumeParser
from app.parsers.nlp_extractor import NLPExtractor

logger = get_logger(__name__)


class DOCXResumeParser(BaseResumeParser):
    PARSER_VERSION = "1.0.0"

    def __init__(self) -> None:
        self._extractor = NLPExtractor()

    # ── Public API ─────────────────────────────────────────────────────────────

    def extract_text(self, content: bytes) -> str:
        """Extract all text from DOCX preserving paragraph boundaries."""
        import io

        try:
            from docx import Document
            from docx.oxml.ns import qn

            doc = Document(io.BytesIO(content))
            parts: list[str] = []

            # Main body paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    parts.append(text)

            # Tables (common in two-column resumes)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        parts.append(row_text)

            return "\n".join(parts)

        except Exception as exc:
            raise ParseError(
                "Failed to extract text from DOCX.",
                detail=str(exc),
            ) from exc

    def parse(self, content: bytes) -> ParsedResume:
        """
        DOCX parsing with style-awareness:
          - Heading styles → section headers (higher confidence than regex)
          - Bold runs → often company names / job titles
          - List paragraphs → bullet points
        """
        import io

        from docx import Document

        try:
            doc = Document(io.BytesIO(content))
        except Exception as exc:
            raise ParseError("Could not open DOCX document.", detail=str(exc)) from exc

        raw_text = self.extract_text(content)

        # Attempt style-aware parsing first
        style_parsed = self._parse_with_styles(doc)
        if style_parsed:
            sections_override = style_parsed
        else:
            sections_override = None

        sections = self._extractor.split_sections(raw_text)

        # If style parsing found better section content, prefer it
        if sections_override:
            if sections_override.get("experience"):
                sections.experience_lines = sections_override["experience"]
            if sections_override.get("skills"):
                sections.skills_lines = sections_override["skills"]
            if sections_override.get("education"):
                sections.education_lines = sections_override["education"]

        contact = self._extractor.extract_contact(raw_text)
        skills = self._extractor.extract_skills(sections.skills_lines)
        experience = self._extractor.extract_experience(sections.experience_lines)
        education = self._extractor.extract_education(sections.education_lines)
        projects = self._extractor.extract_projects(sections.projects_lines)
        achievements = self._extractor.extract_achievements(sections.achievements_lines)
        certifications = self._extractor.extract_achievements(sections.certifications_lines)
        ats_keywords = self._extractor.extract_ats_keywords(raw_text, skills)

        total_years = self._extractor.estimate_total_experience(experience)
        seniority = self._extractor.infer_seniority(
            total_years, [exp.title for exp in experience]
        )

        confidence = self._compute_confidence(contact, skills, experience, education)

        return ParsedResume(
            contact=contact,
            summary=sections.summary or None,
            skills=skills,
            work_experience=experience,
            education=education,
            projects=projects,
            achievements=achievements,
            certifications=certifications,
            ats_keywords=ats_keywords,
            total_years_experience=total_years,
            seniority_level=seniority,
            raw_text_length=len(raw_text),
            parser_version=self.PARSER_VERSION,
            confidence_score=confidence,
        )

    # ── Style-aware extraction ─────────────────────────────────────────────────

    def _parse_with_styles(self, doc) -> dict[str, list[str]] | None:
        """
        Use paragraph style names to identify sections more reliably.
        Returns a dict of section_name → [lines] or None if styles aren't useful.
        """
        try:
            from app.parsers.nlp_extractor import _SECTION_HEADERS

            result: dict[str, list[str]] = {}
            current_section: str | None = None
            heading_found = False

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                style_name = para.style.name.lower() if para.style else ""
                is_heading = "heading" in style_name or para.runs and all(
                    r.bold for r in para.runs if r.text.strip()
                )

                if is_heading:
                    heading_found = True
                    for section_name, pattern in _SECTION_HEADERS.items():
                        if pattern.match(text):
                            current_section = section_name
                            break
                    else:
                        current_section = None
                elif current_section:
                    result.setdefault(current_section, []).append(text)

            return result if heading_found and result else None

        except Exception:
            return None

    # ── Confidence ────────────────────────────────────────────────────────────

    def _compute_confidence(self, contact, skills, experience, education) -> float:
        score = 0.2  # DOCX gets base boost — structure is more reliable
        if contact.email:
            score += 0.15
        if contact.name:
            score += 0.1
        if skills:
            score += min(0.2, len(skills) * 0.02)
        if experience:
            score += min(0.2, len(experience) * 0.07)
        if education:
            score += 0.1
        return round(min(score, 1.0), 3)
